from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "entregables"
OUT.mkdir(parents=True, exist_ok=True)


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
GRAY = "F2F4F7"
GREEN = "D9EAD3"
YELLOW = "FFF2CC"
RED = "F4CCCC"


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 121)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.name = "Arial"
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def read_md_sections(name):
    path = DOCS / name
    return path.read_text(encoding="utf-8").splitlines()


def doc_from_markdown(md_name, out_name, title, subtitle):
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, title, subtitle)
    pending_table = []
    for line in read_md_sections(md_name):
        text = line.strip()
        if not text:
            if pending_table:
                emit_markdown_table(doc, pending_table)
                pending_table = []
            continue
        if text.startswith("|"):
            pending_table.append(text)
            continue
        if pending_table:
            emit_markdown_table(doc, pending_table)
            pending_table = []
        if text.startswith("# "):
            doc.add_heading(text[2:], level=1)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=2)
        elif text.startswith("### "):
            doc.add_heading(text[4:], level=3)
        elif text.startswith("- "):
            doc.add_paragraph(text[2:], style="List Bullet")
        elif text.startswith("```"):
            continue
        else:
            doc.add_paragraph(text.replace("`", ""))
    if pending_table:
        emit_markdown_table(doc, pending_table)
    doc.save(OUT / out_name)


def emit_markdown_table(doc, lines):
    rows = []
    for line in lines:
        parts = [part.strip() for part in line.strip("|").split("|")]
        if all(set(part) <= {"-", ":"} for part in parts):
            continue
        rows.append(parts)
    if len(rows) < 2:
        return
    add_table(doc, rows[0], rows[1:])


def build_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "Casos de Prueba"
    headers = [
        "ID", "Modulo", "Caso", "Tipo", "Prioridad", "Precondicion",
        "Pasos", "Resultado Esperado", "Estado", "Evidencia"
    ]
    rows = [
        ["CP-001", "Setup", "Crear admin inicial", "Sistema", "Alta", "App desplegada", "GET /setup-admin", "Usuario creado o existente", "Preparado", "target/postman"],
        ["CP-002", "Login", "Login admin", "E2E", "Alta", "Admin existe", "Ingresar credenciales", "Dashboard visible", "Preparado", "target/playwright"],
        ["CP-003", "Producto", "Validar campos requeridos", "Postman/Playwright", "Alta", "Sesion admin", "Abrir formulario", "Campos required presentes", "Preparado", "target/postman"],
        ["CP-004", "Producto", "Crear producto", "E2E", "Alta", "Sesion admin", "Completar y guardar", "Producto en listado", "Preparado", "target/playwright"],
        ["CP-005", "Tiendas", "Listar tiendas", "Sistema", "Media", "Sesion admin", "Abrir /tiendas", "Listado sin error", "Preparado", "target/failsafe-reports"],
        ["CP-006", "Pedidos", "Crear pedido", "Integracion", "Alta", "Tienda existe", "Crear pedido", "Pedido persistido", "Preparado", "target/surefire-reports"],
        ["CP-007", "Pedidos Especiales", "Guardar con imagen", "Unitaria", "Media", "Tienda existe", "Adjuntar imagen", "Imagen procesada", "Preparado", "target/surefire-reports"],
        ["CP-008", "Seguridad", "Acceso protegido sin sesion", "Sistema", "Alta", "Sin sesion", "GET /productos", "Redireccion a login", "Preparado", "target/postman"],
    ]
    ws.append(headers)
    for row in rows:
        ws.append(row)
    style_sheet(ws, "A1:J9", headers=True)
    table = Table(displayName="TablaCasosPrueba", ref="A1:J9")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
    ws.add_table(table)
    ws.freeze_panes = "A2"

    defects = wb.create_sheet("Defectos")
    defect_headers = ["ID", "Fecha", "Modulo", "Severidad", "Descripcion", "Estado", "Responsable", "Evidencia", "Fecha Cierre"]
    defect_rows = [
        ["DEF-001", "", "SonarQube", "Alta", "Registrar bugs/code smells encontrados por Quality Gate", "Pendiente evidencia", "Equipo QA", "SonarQube", ""],
        ["DEF-002", "", "OWASP ZAP", "Media", "Registrar alertas medias/altas encontradas por ZAP", "Pendiente ejecucion", "Equipo QA", "target/security", ""],
    ]
    defects.append(defect_headers)
    for row in defect_rows:
        defects.append(row)
    style_sheet(defects, "A1:I3", headers=True)
    defects.freeze_panes = "A2"

    matrix = wb.create_sheet("Matriz Entregables")
    matrix_headers = ["Item", "Entregable", "Archivo/Evidencia", "Responsable", "Estado"]
    matrix_rows = [
        [1, "Codigo fuente con unitarias", "src/test/java", "Dev", "Preparado"],
        [2, "Postman integrales", "postman/EspigaPedidos.postman_collection.json", "QA", "Preparado"],
        [3, "SonarQube cobertura >=80%", "target/site/jacoco + Sonar", "QA", "Pendiente captura"],
        [4, "Informe SonarQube Nivel A", "SonarQube export", "QA", "Pendiente exportar"],
        [5, "Manual CI/CD", "docs/MANUAL_CONFIGURACION_CICD.md", "DevOps", "Preparado"],
        [6, "Informe pruebas sistema", "docs/entregables/Informe_Pruebas_Sistema_E2E.docx", "QA", "Preparado"],
        [7, "Excel casos/defectos", "docs/entregables/Gestion_Casos_Prueba_Defectos.xlsx", "QA", "Preparado"],
        [8, "Informe seguridad", "docs/INFORME_SEGURIDAD.md", "QA", "Preparado"],
        [8.1, "Informe rendimiento", "docs/INFORME_RENDIMIENTO.md", "QA", "Preparado"],
        [9, "Manual usuario", "docs/entregables/Manual_Usuario_EspigaPedidos.docx", "QA", "Preparado"],
        [10, "App desplegada URL", "APP_BASE_URL", "DevOps", "Preparado"],
        [11, "Ambiente pruebas", "docker-compose.yml", "DevOps", "Preparado"],
    ]
    matrix.append(matrix_headers)
    for row in matrix_rows:
        matrix.append(row)
    style_sheet(matrix, "A1:E13", headers=True)
    matrix.freeze_panes = "A2"

    sonar = wb.create_sheet("SonarQube")
    sonar_headers = ["Metrica", "Criterio", "Resultado local", "Evidencia", "Estado"]
    sonar_rows = [
        ["Bugs", "0", "Pendiente dashboard", "SonarQube", "Pendiente exportar"],
        ["Vulnerabilidades", "0", "Pendiente dashboard", "SonarQube", "Pendiente exportar"],
        ["Code Smells criticos", "0", "Pendiente dashboard", "SonarQube", "Pendiente exportar"],
        ["Maintainability Rating", "A", "Pendiente dashboard", "SonarQube", "Pendiente exportar"],
        ["Reliability Rating", "A", "Pendiente dashboard", "SonarQube", "Pendiente exportar"],
        ["Security Rating", "A", "Pendiente dashboard", "SonarQube", "Pendiente exportar"],
        ["Line Coverage", ">= 80%", "100.00%", "target/site/jacoco/index.html", "Superado"],
        ["Instruction Coverage", ">= 80%", "100.00%", "target/site/jacoco/jacoco.xml", "Superado"],
        ["Branch Coverage", "Informativo", "93.75%", "target/site/jacoco/index.html", "Superado"],
    ]
    sonar.append(sonar_headers)
    for row in sonar_rows:
        sonar.append(row)
    style_sheet(sonar, "A1:E10", headers=True)
    sonar.freeze_panes = "A2"

    for sheet in wb.worksheets:
        for column_cells in sheet.columns:
            width = min(max(len(str(cell.value or "")) for cell in column_cells) + 3, 55)
            sheet.column_dimensions[column_cells[0].column_letter].width = width
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    wb.save(OUT / "Gestion_Casos_Prueba_Defectos.xlsx")


def style_sheet(ws, ref, headers=False):
    thin = Side(style="thin", color="D9E2F3")
    for row in ws[ref]:
        for cell in row:
            cell.border = Border(bottom=thin)
            cell.font = Font(name="Arial", size=10)
    if headers:
        for cell in ws[1]:
            cell.font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=BLUE)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def main():
    doc_from_markdown(
        "MANUAL_USUARIO.md",
        "Manual_Usuario_EspigaPedidos.docx",
        "Manual de Usuario",
        "Sistema EspigaPedidos"
    )
    doc_from_markdown(
        "MANUAL_CONFIGURACION_CICD.md",
        "Manual_Configuracion_CICD.docx",
        "Manual de Configuracion CI/CD",
        "Jenkins, SonarQube, Docker, Postman, Playwright, ZAP y k6"
    )
    doc_from_markdown(
        "INFORME_SONARQUBE.md",
        "Informe_SonarQube_Nivel_A.docx",
        "Informe SonarQube - Nivel A",
        "Cobertura, Quality Gate, bugs y code smells"
    )
    doc_from_markdown(
        "INFORME_PRUEBAS_SISTEMA.md",
        "Informe_Pruebas_Sistema_E2E.docx",
        "Informe de Pruebas de Sistema y EndToEnd",
        "Unitarias, integrales, Selenium, Playwright y Postman"
    )
    doc_from_markdown(
        "INFORME_SEGURIDAD.md",
        "Informe_Pruebas_Seguridad.docx",
        "Informe de Pruebas de Seguridad",
        "OWASP ZAP y SonarQube"
    )
    doc_from_markdown(
        "INFORME_RENDIMIENTO.md",
        "Informe_Pruebas_Rendimiento.docx",
        "Informe de Pruebas de Rendimiento",
        "k6 y caracteristicas del servidor"
    )
    build_workbook()


if __name__ == "__main__":
    main()
